"""LLM-driven proofreader for translated .docx documents.

Compares a target .docx against an optional source .docx and rewrites the target
to be a better translation. Loops the LLM to convergence (text-equality stop).

Usage:
    python src/proofreader.py path/to/translated.docx --source path/to/original.docx
    python src/proofreader.py path/to/translated.docx --max-iterations 1
"""

import argparse
import io
import math
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from src.helpers import gpt_41_chat

CHUNK_CHAR_BUDGET = 40_000
SOURCE_CHAR_BUDGET = 60_000
SOURCE_PAD_PARAGRAPHS = 5
MAX_ABSOLUTE_CHARS = 5_000_000
DEFAULT_MAX_ITERATIONS = 3
MAX_OUTPUT_TOKENS = 16_000

_HYPERLINK = qn("w:hyperlink")
_FLD_SIMPLE = qn("w:fldSimple")
_FLD_CHAR = qn("w:fldChar")
_DRAWING = qn("w:drawing")
_RUN = qn("w:r")
_P_PR = qn("w:pPr")

SYSTEM_PROMPT_TRANSLATION = """\
You are a professional proofreader specialising in English-French translation quality.

You will receive TARGET paragraphs (the translated document) and SOURCE paragraphs (the original).
Use SOURCE as reference context to verify the translation is faithful, accurate, and natural.
Do NOT assume one-to-one paragraph mapping between SOURCE and TARGET — translation routinely
merges or splits paragraphs. Only TARGET paragraphs are rewritten.

Your tasks:
1. Fix mistranslations, false cognates, and incorrect word choice.
2. Fix awkward or unnatural phrasing caused by literal translation.
3. Fix grammar, spelling, and punctuation errors.
4. Apply correct French punctuation spacing rules (space before : ; ! ?) if the document is in French.
5. Preserve TARGET paragraph meaning. DO NOT add, remove, or reorder paragraphs.
6. Do not insert newlines within a paragraph — each [N] block must be a single line.

Respond in EXACTLY this format — no extra text before or after:

CHANGED
[N] <full corrected text of TARGET paragraph N>
[M] <full corrected text of TARGET paragraph M>
...
CHANGES
1. [N] "<original phrase>" -> "<corrected phrase>" - <brief reason>
2. [M] "<original phrase>" -> "<corrected phrase>" - <brief reason>
...

Only include paragraphs in CHANGED that actually need rewriting. List them in ascending order
of N. Do NOT echo unchanged paragraphs. The CHANGED block contains the FULL corrected text of
each listed paragraph (not a diff).

If NO TARGET paragraphs need any changes, write exactly:
CHANGED
NONE
CHANGES
NONE
"""

SYSTEM_PROMPT_TARGET_ONLY = """\
You are a professional proofreader.

You will receive paragraphs of a document, each prefixed with [N].

Your tasks:
1. Fix grammar, spelling, and punctuation errors.
2. Fix awkward or unnatural phrasing.
3. Apply correct French punctuation spacing rules (space before : ; ! ?) if the document is in French.
4. Preserve paragraph meaning. DO NOT add, remove, or reorder paragraphs.
5. Do not insert newlines within a paragraph — each [N] block must be a single line.

Respond in EXACTLY this format — no extra text before or after:

CHANGED
[N] <full corrected text of paragraph N>
[M] <full corrected text of paragraph M>
...
CHANGES
1. [N] "<original phrase>" -> "<corrected phrase>" - <brief reason>
2. [M] "<original phrase>" -> "<corrected phrase>" - <brief reason>
...

Only include paragraphs in CHANGED that actually need rewriting. List them in ascending order
of N. Do NOT echo unchanged paragraphs. The CHANGED block contains the FULL corrected text of
each listed paragraph (not a diff).

If NO paragraphs need any changes, write exactly:
CHANGED
NONE
CHANGES
NONE
"""


def read_docx_paragraphs(source):
    """Return paragraph text from a .docx, preserving order and empty paragraphs.

    `source` may be a path, a file-like object, or bytes.
    """
    if isinstance(source, (bytes, bytearray)):
        source = io.BytesIO(source)
    doc = Document(source)
    return [p.text for p in doc.paragraphs]


def build_user_message(target_paragraphs, source_paragraphs=None):
    """Build the single-string user message for gpt_41_chat."""
    
    def render_block(label, paragraphs, prefix):
        lines = [label]
        for i, text in enumerate(paragraphs, start=1):
            single_line = text.replace("\n", " ").replace("\r", " ")
            lines.append(f"[{prefix}{i}] {single_line}")
        return "\n".join(lines)
    
    parts = []
    if source_paragraphs is not None:
        parts.append(render_block("SOURCE", source_paragraphs, "s"))
        parts.append("")
    parts.append(render_block("TARGET", target_paragraphs, ""))
    return "\n".join(parts)


_PARA_LINE_RE = re.compile(r"^\[(\d+)\]\s?(.*)$")


def parse_response(response, expected_paragraph_count):
    """Parse the model's CHANGED / CHANGES diff response.

    Returns:
        (changed_by_index: dict[int, str], changes: list[str])

    Indices in changed_by_index are 1-based. Raises ValueError on malformed
    output or out-of-range paragraph indices.
    """
    text = response.strip()
    if "CHANGED" not in text or "CHANGES" not in text:
        raise ValueError(
            "Response missing CHANGED or CHANGES section.\n--- raw response ---\n" + response
        )
    
    changed_part, _, changes_part = text.partition("CHANGES")
    changed_part = changed_part.split("CHANGED", 1)[1].strip()
    changes_part = changes_part.strip()
    
    changed_by_index = {}
    if not changed_part.upper().startswith("NONE"):
        for line in changed_part.splitlines():
            line = line.rstrip()
            if not line:
                continue
            match = _PARA_LINE_RE.match(line)
            if not match:
                continue
            idx = int(match.group(1))
            if idx < 1 or idx > expected_paragraph_count:
                raise ValueError(
                    f"CHANGED block references paragraph [{idx}] outside valid range "
                    f"1..{expected_paragraph_count}."
                )
            changed_by_index[idx] = match.group(2)
    
    if changes_part.upper().startswith("NONE"):
        changes = []
    else:
        changes = [line.strip() for line in changes_part.splitlines() if line.strip()]
    
    return changed_by_index, changes


def _paragraph_is_complex(paragraph):
    """True if the paragraph has hyperlinks, fields, or embedded drawings."""
    p_elem = paragraph._p
    for child in p_elem:
        tag = child.tag
        if tag in (_HYPERLINK, _FLD_SIMPLE, _FLD_CHAR):
            return True
        if tag == _RUN and child.find(f".//{_DRAWING}") is not None:
            return True
        if tag == _RUN and child.find(f".//{_FLD_CHAR}") is not None:
            return True
    return False


def _rewrite_simple_paragraph(paragraph, new_text):
    """Replace the text of a simple paragraph while preserving the first run's
    style. Clears subsequent runs but leaves their elements in place."""
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(new_text)
        return
    runs[0].text = new_text
    for run in runs[1:]:
        run.text = ""


def apply_corrections(source, corrected_paragraphs, output):
    """Open the source docx, apply corrections, write to output.

    `source` and `output` may be a path or file-like object. Returns a list of
    warning strings for paragraphs we refused to rewrite.
    """
    if isinstance(source, (bytes, bytearray)):
        source = io.BytesIO(source)
    doc = Document(source)
    warnings = []
    
    paragraphs = doc.paragraphs
    if len(paragraphs) != len(corrected_paragraphs):
        raise ValueError(
            f"Corrected paragraph count {len(corrected_paragraphs)} does not match "
            f"document paragraph count {len(paragraphs)}."
        )
    
    for i, (paragraph, new_text) in enumerate(zip(paragraphs, corrected_paragraphs), start=1):
        original = paragraph.text
        if new_text == original:
            continue
        if _paragraph_is_complex(paragraph):
            warnings.append(
                f"Paragraph {i} skipped (contains hyperlink, field, or image); please correct manually. "
                f'Suggested: "{new_text}"'
            )
            continue
        _rewrite_simple_paragraph(paragraph, new_text)
    
    doc.save(output)
    return warnings


def _render_changes_text(iteration_blocks, warnings):
    lines = []
    any_changes = False
    for header, changes in iteration_blocks:
        lines.append(f"=== {header} ===")
        if not changes:
            lines.append("NONE")
        else:
            any_changes = True
            for j, change in enumerate(changes, start=1):
                lines.append(f"{j}. {change}")
        lines.append("")
    
    if warnings:
        lines.append("=== Warnings ===")
        for warning in warnings:
            lines.append(f"- {warning}")
        lines.append("")
    
    if not any_changes and not warnings:
        return "No changes.\n"
    
    return "\n".join(lines).rstrip() + "\n"


def _total_chars(target_paragraphs, source_paragraphs):
    total = sum(len(p) for p in target_paragraphs)
    if source_paragraphs is not None:
        total += sum(len(p) for p in source_paragraphs)
    return total


def _chunk_target_paragraphs(target_paragraphs):
    chunks = []
    current = []
    current_chars = 0
    start_idx = 0
    
    for i, para in enumerate(target_paragraphs):
        para_len = len(para)
        if current and current_chars + para_len > CHUNK_CHAR_BUDGET:
            chunks.append((start_idx, current))
            current = []
            current_chars = 0
            start_idx = i
        current.append(para)
        current_chars += para_len
    
    if current:
        chunks.append((start_idx, current))
    
    if not chunks:
        chunks.append((0, list(target_paragraphs)))
    
    return chunks


def _slice_source_for_chunk(source_paragraphs, target_start, target_end, target_total):
    if source_paragraphs is None:
        return None
    m = len(source_paragraphs)
    if m == 0 or target_total == 0:
        return list(source_paragraphs)
    
    anchor_start = target_start * m / target_total
    anchor_end = target_end * m / target_total
    s_start = max(0, math.floor(anchor_start) - SOURCE_PAD_PARAGRAPHS)
    s_end = min(m, math.ceil(anchor_end) + SOURCE_PAD_PARAGRAPHS)
    
    while (s_end - s_start) > 1 and sum(len(p) for p in source_paragraphs[s_start:s_end]) > SOURCE_CHAR_BUDGET:
        left_pad = anchor_start - s_start
        right_pad = s_end - anchor_end
        if right_pad >= left_pad:
            s_end -= 1
        else:
            s_start += 1
    
    return source_paragraphs[s_start:s_end]


def _rewrite_change_lines_global(changes, global_offset):
    if global_offset == 0:
        return list(changes)
    
    rewritten = []
    for line in changes:
        match = re.search(r"\[(\d+)\]", line)
        if not match:
            rewritten.append(line)
            continue
        local_idx = int(match.group(1))
        global_idx = local_idx + global_offset
        rewritten.append(line.replace(f"[{local_idx}]", f"[{global_idx}]", 1))
    return rewritten


def _proofread_chunk(chunk_paragraphs, source_slice, system_prompt, max_iterations):
    current = list(chunk_paragraphs)
    changes_per_iteration = []
    previous_changes = None
    
    for _ in range(max_iterations):
        user_msg = build_user_message(current, source_slice)
        response = gpt_41_chat(user_msg, system=system_prompt, max_output_tokens=MAX_OUTPUT_TOKENS)
        changed, changes = parse_response(response, expected_paragraph_count=len(current))
        
        if not changed:
            break
        
        oscillating = previous_changes is not None and changes == previous_changes
        for idx, new_text in changed.items():
            current[idx - 1] = new_text
        changes_per_iteration.append(changes)
        previous_changes = changes
        if oscillating:
            break
    
    return current, changes_per_iteration


def proofread_bytes(
        target_bytes,
        source_bytes=None,
        target_filename="document.docx",
        max_iterations=DEFAULT_MAX_ITERATIONS,
        progress_callback=None,
):
    """Proofread target_bytes against optional source_bytes.

    Returns (corrected_docx_bytes, changes_text).
    """
    target_paragraphs = read_docx_paragraphs(target_bytes)
    source_paragraphs = read_docx_paragraphs(source_bytes) if source_bytes else None
    
    total = _total_chars(target_paragraphs, source_paragraphs)
    if total > MAX_ABSOLUTE_CHARS:
        raise ValueError(
            f"Document too large to proofread ({total:,} chars; "
            f"limit {MAX_ABSOLUTE_CHARS:,})."
        )
    
    system_prompt = (
        SYSTEM_PROMPT_TRANSLATION if source_paragraphs is not None else SYSTEM_PROMPT_TARGET_ONLY
    )
    
    chunks = _chunk_target_paragraphs(target_paragraphs)
    total_chunks = len(chunks)
    current = list(target_paragraphs)
    iteration_blocks = []
    
    for chunk_idx, (global_start, chunk_paragraphs) in enumerate(chunks):
        chunk_end = global_start + len(chunk_paragraphs)
        source_slice = _slice_source_for_chunk(
            source_paragraphs, global_start, chunk_end, len(target_paragraphs)
        )
        updated_chunk, chunk_changes_per_iter = _proofread_chunk(
            chunk_paragraphs, source_slice, system_prompt, max_iterations
        )
        for k, new_text in enumerate(updated_chunk):
            current[global_start + k] = new_text
        
        for iter_idx, changes in enumerate(chunk_changes_per_iter, start=1):
            global_changes = _rewrite_change_lines_global(changes, global_start)
            if total_chunks == 1:
                header = f"Iteration {iter_idx}"
            else:
                header = f"Chunk {chunk_idx + 1}/{total_chunks}, Iteration {iter_idx}"
            iteration_blocks.append((header, global_changes))
        
        if progress_callback is not None:
            progress_callback(chunk_idx + 1, total_chunks)
    
    output_buffer = io.BytesIO()
    warnings = apply_corrections(target_bytes, current, output_buffer)
    output_buffer.seek(0)
    
    changes_text = _render_changes_text(iteration_blocks, warnings)
    return output_buffer.getvalue(), changes_text


def proofread(target_path, source_path=None, max_iterations=DEFAULT_MAX_ITERATIONS, output_dir=None):
    """File-based entrypoint. Writes <stem>_proofread.docx and <stem>_changes.txt."""
    target_path = Path(target_path)
    if not target_path.exists():
        raise FileNotFoundError(f"Target document not found: {target_path}")
    if target_path.suffix.lower() != ".docx":
        raise ValueError(f"Target must be a .docx file: {target_path}")
    
    target_bytes = target_path.read_bytes()
    
    source_bytes = None
    if source_path is not None:
        source_path = Path(source_path)
        if not source_path.exists():
            raise FileNotFoundError(f"Source document not found: {source_path}")
        if source_path.suffix.lower() != ".docx":
            raise ValueError(f"Source must be a .docx file: {source_path}")
        source_bytes = source_path.read_bytes()
    
    out_docx_bytes, changes_text = proofread_bytes(
        target_bytes=target_bytes,
        source_bytes=source_bytes,
        target_filename=target_path.name,
        max_iterations=max_iterations,
    )
    
    output_dir = Path(output_dir) if output_dir is not None else target_path.parent
    output_dir.mkdir(parents=True, exist_ok=True)
    stem = target_path.stem
    out_docx_path = output_dir / f"{stem}_proofread.docx"
    out_changes_path = output_dir / f"{stem}_changes.txt"
    
    out_docx_path.write_bytes(out_docx_bytes)
    out_changes_path.write_text(changes_text, encoding="utf-8")
    
    print(f"Proofread:  {out_docx_path}")
    print(f"Changes:    {out_changes_path}")


def main():
    parser = argparse.ArgumentParser(description="Proofread a translated .docx against an optional source.")
    parser.add_argument("target", help="Path to the translated .docx to proofread.")
    parser.add_argument("--source", default=None, help="Path to the original .docx (optional).")
    parser.add_argument(
        "--max-iterations",
        type=int,
        default=DEFAULT_MAX_ITERATIONS,
        help=f"Maximum LLM passes before stopping (default: {DEFAULT_MAX_ITERATIONS}).",
    )
    parser.add_argument("--output-dir", default=None, help="Where to write outputs (defaults to target's directory).")
    args = parser.parse_args()
    
    try:
        proofread(
            target_path=args.target,
            source_path=args.source,
            max_iterations=args.max_iterations,
            output_dir=args.output_dir,
        )
    except (FileNotFoundError, ValueError) as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
