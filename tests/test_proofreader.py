import io

import pytest
from docx import Document
from docx.oxml import OxmlElement

from src import proofreader
from src.proofreader import (
    SYSTEM_PROMPT_TARGET_ONLY,
    SYSTEM_PROMPT_TRANSLATION,
    _chunk_target_paragraphs,
    _paragraph_is_complex,
    _render_changes_text,
    _rewrite_change_lines_global,
    _rewrite_simple_paragraph,
    _slice_source_for_chunk,
    _total_chars,
    apply_corrections,
    build_user_message,
    parse_response,
    proofread_bytes,
    read_docx_paragraphs,
)


def _build_complex_docx():
    # A paragraph carrying a hyperlink element — the kind apply_corrections
    # refuses to rewrite.
    doc = Document()
    paragraph = doc.add_paragraph("complex paragraph")
    paragraph._p.append(OxmlElement("w:hyperlink"))
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


class FakeLLM:
    def __init__(self, response):
        self.response = response
        self.calls = []
    
    def __call__(self, msg, system=None, max_output_tokens=None, **kwargs):
        self.calls.append({"msg": msg, "system": system})
        return self.response


class TestReadDocxParagraphs:
    def test_reads_paragraphs_from_bytes(self, make_docx):
        data = make_docx(["first", "second"])
        assert read_docx_paragraphs(data) == ["first", "second"]
    
    def test_reads_paragraphs_from_file_like(self, make_docx):
        data = make_docx(["first", "second"])
        assert read_docx_paragraphs(io.BytesIO(data)) == ["first", "second"]
    
    def test_reads_paragraphs_from_path(self, make_docx, tmp_path):
        path = tmp_path / "doc.docx"
        path.write_bytes(make_docx(["alpha", "beta"]))
        assert read_docx_paragraphs(path) == ["alpha", "beta"]
    
    def test_preserves_empty_paragraphs(self, make_docx):
        data = make_docx(["a", "", "b"])
        assert read_docx_paragraphs(data) == ["a", "", "b"]


class TestBuildUserMessage:
    def test_target_only(self):
        message = build_user_message(["one", "two"])
        assert message == "TARGET\n[1] one\n[2] two"
    
    def test_with_source(self):
        message = build_user_message(["one"], ["src one", "src two"])
        assert message == "SOURCE\n[s1] src one\n[s2] src two\n\nTARGET\n[1] one"
    
    def test_newlines_collapsed_to_spaces(self):
        message = build_user_message(["line one\nline two"])
        assert message == "TARGET\n[1] line one line two"


class TestParseResponse:
    def test_parses_changed_and_changes(self):
        response = (
            "CHANGED\n"
            "[1] First corrected.\n"
            "[3] Third corrected.\n"
            "CHANGES\n"
            '1. [1] "first" -> "First corrected" - capitalisation\n'
            '2. [3] "third" -> "Third corrected" - capitalisation\n'
        )
        changed, changes = parse_response(response, expected_paragraph_count=3)
        assert changed == {1: "First corrected.", 3: "Third corrected."}
        assert len(changes) == 2
    
    def test_none_response_yields_empty_results(self):
        response = "CHANGED\nNONE\nCHANGES\nNONE\n"
        assert parse_response(response, expected_paragraph_count=3) == ({}, [])
    
    def test_missing_changes_section_raises(self):
        with pytest.raises(ValueError):
            parse_response("CHANGED\n[1] x", expected_paragraph_count=1)
    
    def test_missing_changed_section_raises(self):
        with pytest.raises(ValueError):
            parse_response("CHANGES\nNONE", expected_paragraph_count=1)
    
    def test_out_of_range_index_raises(self):
        response = "CHANGED\n[5] out of range\nCHANGES\nNONE\n"
        with pytest.raises(ValueError):
            parse_response(response, expected_paragraph_count=3)
    
    def test_unparseable_lines_are_skipped(self):
        response = (
            "CHANGED\n"
            "[1] kept\n"
            "stray line with no index\n"
            "[2] also kept\n"
            "CHANGES\nNONE\n"
        )
        changed, _ = parse_response(response, expected_paragraph_count=2)
        assert changed == {1: "kept", 2: "also kept"}


class TestParagraphIsComplex:
    def test_plain_paragraph_is_simple(self):
        doc = Document()
        paragraph = doc.add_paragraph("just text")
        assert _paragraph_is_complex(paragraph) is False
    
    def test_paragraph_with_hyperlink_is_complex(self):
        doc = Document()
        paragraph = doc.add_paragraph("link text")
        paragraph._p.append(OxmlElement("w:hyperlink"))
        assert _paragraph_is_complex(paragraph) is True


class TestRewriteSimpleParagraph:
    def test_rewrites_first_run_and_clears_the_rest(self):
        doc = Document()
        paragraph = doc.add_paragraph()
        paragraph.add_run("foo")
        paragraph.add_run("bar")
        _rewrite_simple_paragraph(paragraph, "replacement")
        assert paragraph.runs[0].text == "replacement"
        assert paragraph.runs[1].text == ""
        assert paragraph.text == "replacement"
    
    def test_adds_a_run_when_paragraph_has_none(self):
        doc = Document()
        paragraph = doc.add_paragraph()
        _rewrite_simple_paragraph(paragraph, "added text")
        assert paragraph.text == "added text"


class TestApplyCorrections:
    def test_applies_changed_paragraphs(self, make_docx):
        data = make_docx(["one", "two"])
        output = io.BytesIO()
        warnings = apply_corrections(data, ["one", "two corrected"], output)
        assert warnings == []
        output.seek(0)
        assert read_docx_paragraphs(output) == ["one", "two corrected"]
    
    def test_paragraph_count_mismatch_raises(self, make_docx):
        data = make_docx(["one", "two"])
        with pytest.raises(ValueError):
            apply_corrections(data, ["only one"], io.BytesIO())
    
    def test_complex_paragraph_is_skipped_with_a_warning(self):
        data = _build_complex_docx()
        output = io.BytesIO()
        warnings = apply_corrections(data, ["a different value"], output)
        assert len(warnings) == 1
        assert "skipped" in warnings[0]
        output.seek(0)
        assert read_docx_paragraphs(output) != ["a different value"]


class TestRenderChangesText:
    def test_no_blocks_and_no_warnings(self):
        assert _render_changes_text([], []) == "No changes.\n"
    
    def test_renders_iteration_changes(self):
        rendered = _render_changes_text([("Iteration 1", ["a change"])], [])
        assert "=== Iteration 1 ===" in rendered
        assert "1. a change" in rendered
    
    def test_renders_warnings_section(self):
        rendered = _render_changes_text([], ["something was skipped"])
        assert "=== Warnings ===" in rendered
        assert "- something was skipped" in rendered
    
    def test_iteration_with_no_changes_reports_no_changes(self):
        assert _render_changes_text([("Iteration 1", [])], []) == "No changes.\n"


class TestTotalChars:
    def test_target_only(self):
        assert _total_chars(["ab", "cde"], None) == 5
    
    def test_target_and_source(self):
        assert _total_chars(["ab"], ["xyz"]) == 5


class TestChunkTargetParagraphs:
    def test_small_input_is_a_single_chunk(self):
        chunks = _chunk_target_paragraphs(["a", "b"])
        assert chunks == [(0, ["a", "b"])]
    
    def test_empty_input_yields_one_empty_chunk(self):
        assert _chunk_target_paragraphs([]) == [(0, [])]
    
    def test_large_input_splits_into_multiple_chunks(self):
        paragraphs = ["x" * 15_000] * 4
        chunks = _chunk_target_paragraphs(paragraphs)
        assert len(chunks) == 2
        assert chunks[0][0] == 0
        assert chunks[1][0] == 2


class TestSliceSourceForChunk:
    def test_none_source_returns_none(self):
        assert _slice_source_for_chunk(None, 0, 5, 10) is None
    
    def test_zero_target_total_returns_full_source(self):
        source = ["a", "b"]
        assert _slice_source_for_chunk(source, 0, 1, 0) == source
    
    def test_slices_around_the_chunk_anchor(self):
        source = [f"p{i}" for i in range(30)]
        assert _slice_source_for_chunk(source, 10, 20, 30) == source[5:25]


class TestRewriteChangeLinesGlobal:
    def test_zero_offset_leaves_lines_unchanged(self):
        lines = ['1. [2] "a" -> "b" - reason']
        assert _rewrite_change_lines_global(lines, 0) == lines
    
    def test_offset_shifts_the_paragraph_index(self):
        lines = ['1. [2] "a" -> "b" - reason']
        assert _rewrite_change_lines_global(lines, 5) == ['1. [7] "a" -> "b" - reason']
    
    def test_lines_without_an_index_are_left_alone(self):
        lines = ["no index in this line"]
        assert _rewrite_change_lines_global(lines, 5) == lines


class TestProofreadBytes:
    def test_no_changes_returns_original_document(self, make_docx, monkeypatch):
        fake = FakeLLM("CHANGED\nNONE\nCHANGES\nNONE\n")
        monkeypatch.setattr(proofreader, "gpt_41_chat", fake)
        
        data = make_docx(["Hello world.", "Second."])
        out_bytes, changes_text = proofread_bytes(data, max_iterations=1)
        
        assert read_docx_paragraphs(out_bytes) == ["Hello world.", "Second."]
        assert changes_text == "No changes.\n"
    
    def test_applies_a_suggested_change(self, make_docx, monkeypatch):
        response = (
            "CHANGED\n"
            "[1] Hello, world!\n"
            "CHANGES\n"
            '1. [1] "Hello world." -> "Hello, world!" - added comma\n'
        )
        fake = FakeLLM(response)
        monkeypatch.setattr(proofreader, "gpt_41_chat", fake)
        
        data = make_docx(["Hello world.", "Second."])
        out_bytes, changes_text = proofread_bytes(data, max_iterations=1)
        
        assert read_docx_paragraphs(out_bytes) == ["Hello, world!", "Second."]
        assert "Hello, world!" in changes_text
    
    def test_target_only_uses_the_target_only_prompt(self, make_docx, monkeypatch):
        fake = FakeLLM("CHANGED\nNONE\nCHANGES\nNONE\n")
        monkeypatch.setattr(proofreader, "gpt_41_chat", fake)
        
        proofread_bytes(make_docx(["text"]), max_iterations=1)
        assert fake.calls[0]["system"] == SYSTEM_PROMPT_TARGET_ONLY
    
    def test_with_source_uses_the_translation_prompt(self, make_docx, monkeypatch):
        fake = FakeLLM("CHANGED\nNONE\nCHANGES\nNONE\n")
        monkeypatch.setattr(proofreader, "gpt_41_chat", fake)
        
        proofread_bytes(
            make_docx(["text"]),
            source_bytes=make_docx(["texte"]),
            max_iterations=1,
        )
        assert fake.calls[0]["system"] == SYSTEM_PROMPT_TRANSLATION
    
    def test_progress_callback_is_invoked(self, make_docx, monkeypatch):
        fake = FakeLLM("CHANGED\nNONE\nCHANGES\nNONE\n")
        monkeypatch.setattr(proofreader, "gpt_41_chat", fake)
        
        progress = []
        proofread_bytes(
            make_docx(["text"]),
            max_iterations=1,
            progress_callback=lambda done, total: progress.append((done, total)),
        )
        assert progress == [(1, 1)]
    
    def test_oversized_document_is_rejected(self, make_docx, monkeypatch):
        fake = FakeLLM("CHANGED\nNONE\nCHANGES\nNONE\n")
        monkeypatch.setattr(proofreader, "gpt_41_chat", fake)
        monkeypatch.setattr(proofreader, "MAX_ABSOLUTE_CHARS", 5)
        
        with pytest.raises(ValueError):
            proofread_bytes(make_docx(["this is well over five chars"]), max_iterations=1)
