import io
import random

import pytest
from docx import Document
from PIL import Image


# Dict that also supports attribute access, mimicking st.session_state closely
# enough for i18n / tab_guard (which use both `state["k"]` and `state.k`).
class FakeSessionState(dict):
    def __getattr__(self, name):
        if name in self:
            return self[name]
        raise AttributeError(name)
    
    def __setattr__(self, name, value):
        self[name] = value
    
    def __delattr__(self, name):
        del self[name]


@pytest.fixture
def fake_session_state(monkeypatch):
    state = FakeSessionState()
    monkeypatch.setattr("streamlit.session_state", state, raising=False)
    return state


@pytest.fixture
def make_docx():
    def build(paragraphs):
        doc = Document()
        for text in paragraphs:
            doc.add_paragraph(text)
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
    
    return build


@pytest.fixture
def noise_image():
    # Random-noise image: many distinct colours and poor JPEG compressibility,
    # so re-encoding at lower quality / smaller size reliably shrinks it.
    def build(width, height, fmt="JPEG", quality=95, seed=0):
        raw = random.Random(seed).randbytes(width * height * 3)
        img = Image.frombytes("RGB", (width, height), raw)
        buffer = io.BytesIO()
        if fmt == "JPEG":
            img.save(buffer, format=fmt, quality=quality)
        else:
            img.save(buffer, format=fmt)
        return buffer.getvalue()
    
    return build


@pytest.fixture
def make_docx_with_image(noise_image):
    def build(path, width=1600, height=1600):
        image_bytes = noise_image(width, height, fmt="JPEG", quality=95)
        doc = Document()
        doc.add_paragraph("Document with an embedded image.")
        doc.add_picture(io.BytesIO(image_bytes))
        doc.save(str(path))
        return path
    
    return build
