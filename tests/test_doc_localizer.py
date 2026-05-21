import io
import re
from pathlib import Path

import pytest
from docx import Document

from src.doc_localizer import (
    localize_document,
    localize_document_bytes,
    localize_text,
)

NBSP = " "
FIXTURE_DIR = Path(__file__).parent / "fixtures"


def _paragraph(docx_bytes):
    return Document(io.BytesIO(docx_bytes)).paragraphs[0]


class TestNumbersEnglishToFrench:
    @pytest.mark.parametrize(
        "english, french",
        [
            ("1,233.45", f"1{NBSP}233,45"),
            ("1,250", f"1{NBSP}250"),
            ("12,000", f"12{NBSP}000"),
            ("1,234,567.89", f"1{NBSP}234{NBSP}567,89"),
            ("0.0", "0,0"),
            ("123.5", "123,5"),
            ("12.5", "12,5"),
        ],
    )
    def test_number_is_reformatted(self, english, french):
        assert localize_text(english, "fr")[0] == french


class TestNumbersFrenchToEnglish:
    @pytest.mark.parametrize(
        "french, english",
        [
            (f"1{NBSP}233,45", "1,233.45"),
            (f"1{NBSP}250", "1,250"),
            (f"12{NBSP}000", "12,000"),
            (f"1{NBSP}234{NBSP}567,89", "1,234,567.89"),
            ("0,0", "0.0"),
            ("123,5", "123.5"),
            ("12,5", "12.5"),
        ],
    )
    def test_number_is_reformatted(self, french, english):
        assert localize_text(french, "en")[0] == english


class TestBareNumbersAreLeftAlone:
    # Years, page numbers and counts carry no separator and must not be
    # regrouped — "2024" must never become "2 024".
    @pytest.mark.parametrize(
        "text",
        ["2024", "Section 12", "page 5", "1250", "55-65", "1/4", "version 3"],
    )
    def test_unchanged_in_both_directions(self, text):
        assert localize_text(text, "fr")[0] == text
        assert localize_text(text, "en")[0] == text


class TestPercentSpacing:
    @pytest.mark.parametrize(
        "english, french",
        [
            ("50%", f"50{NBSP}%"),
            ("12.5%", f"12,5{NBSP}%"),
            ("1,250%", f"1{NBSP}250{NBSP}%"),
            ("25% - 35%", f"25{NBSP}% - 35{NBSP}%"),
        ],
    )
    def test_english_to_french(self, english, french):
        assert localize_text(english, "fr")[0] == french
    
    @pytest.mark.parametrize(
        "french, english",
        [
            (f"50{NBSP}%", "50%"),
            (f"12,5{NBSP}%", "12.5%"),
            (f"1{NBSP}250{NBSP}%", "1,250%"),
            (f"25{NBSP}% - 35{NBSP}%", "25% - 35%"),
        ],
    )
    def test_french_to_english(self, french, english):
        assert localize_text(french, "en")[0] == english
    
    def test_existing_french_spacing_is_idempotent(self):
        assert localize_text(f"12,5{NBSP}%", "fr")[0] == f"12,5{NBSP}%"


class TestCombinedNumberAndPercent:
    def test_grouped_decimal_percent_english_to_french(self):
        assert localize_text("1,233.45%", "fr")[0] == f"1{NBSP}233,45{NBSP}%"
    
    def test_grouped_decimal_percent_french_to_english(self):
        assert localize_text(f"1{NBSP}233,45{NBSP}%", "en")[0] == "1,233.45%"


class TestPunctuationSpacing:
    @pytest.mark.parametrize(
        "english, french",
        [
            ("Note: here", f"Note{NBSP}: here"),
            ("Wait; then", f"Wait{NBSP}; then"),
            ("Why?", f"Why{NBSP}?"),
            ("Stop!", f"Stop{NBSP}!"),
        ],
    )
    def test_french_adds_nbsp_before_punctuation(self, english, french):
        assert localize_text(english, "fr")[0] == french
    
    @pytest.mark.parametrize(
        "french, english",
        [
            (f"Note{NBSP}: ici", "Note: ici"),
            ("Note : ici", "Note: ici"),
            (f"Pourquoi{NBSP}?", "Pourquoi?"),
            ("Pourquoi ?", "Pourquoi?"),
        ],
    )
    def test_english_removes_space_before_punctuation(self, french, english):
        assert localize_text(french, "en")[0] == english
    
    def test_french_guillemets_get_inner_nbsp(self):
        assert localize_text("«texte»", "fr")[0] == f"«{NBSP}texte{NBSP}»"
    
    @pytest.mark.parametrize(
        "text",
        [
            "see http://example.com now",
            "meet at 10:30 today",
            "ratio 2:1 holds",
        ],
    )
    def test_urls_and_times_are_left_alone(self, text):
        assert localize_text(text, "fr")[0] == text
    
    def test_french_spacing_is_idempotent(self):
        once = localize_text("Note: here", "fr")[0]
        assert localize_text(once, "fr")[0] == once


class TestRoundTrip:
    @pytest.mark.parametrize(
        "english",
        [
            "1,233.45%",
            "0.0",
            "12.5%",
            "Total: 1,250 units; done.",
            "55-65 and 1/4",
        ],
    )
    def test_english_to_french_and_back(self, english):
        french = localize_text(english, "fr")[0]
        assert localize_text(french, "en")[0] == english


class TestChangeCounts:
    def test_counts_a_number_fix(self):
        _, counts = localize_text("0.0", "fr")
        assert counts["numbers"] == 1
    
    def test_counts_a_punctuation_fix(self):
        _, counts = localize_text("Note: here", "fr")
        assert counts["punctuation"] == 1
    
    def test_clean_text_reports_no_changes(self):
        _, counts = localize_text("plain text with no numbers", "fr")
        assert not any(counts.values())
    
    def test_invalid_target_language_raises(self):
        with pytest.raises(ValueError):
            localize_text("anything", "de")


class TestLocalizeDocumentBytes:
    def test_body_paragraph_is_localized(self, make_docx):
        data = make_docx(["The total was 1,233.45 units."])
        out, summary = localize_document_bytes(data, "fr")
        assert _paragraph(out).text == f"The total was 1{NBSP}233,45 units."
        assert summary["fixes"]["numbers"] == 1
        assert summary["text_nodes_scanned"] >= 1
    
    def test_table_cell_is_localized(self):
        doc = Document()
        doc.add_table(rows=1, cols=1).cell(0, 0).paragraphs[0].add_run("0.0")
        buffer = io.BytesIO()
        doc.save(buffer)
        out, _ = localize_document_bytes(buffer.getvalue(), "fr")
        cell = Document(io.BytesIO(out)).tables[0].cell(0, 0)
        assert cell.text == "0,0"
    
    def test_header_is_localized(self):
        doc = Document()
        doc.sections[0].header.paragraphs[0].add_run("Coverage 50%")
        buffer = io.BytesIO()
        doc.save(buffer)
        out, _ = localize_document_bytes(buffer.getvalue(), "fr")
        header = Document(io.BytesIO(out)).sections[0].header
        assert header.paragraphs[0].text == f"Coverage 50{NBSP}%"
    
    def test_clean_document_reports_no_fixes(self, make_docx):
        _, summary = localize_document_bytes(make_docx(["No numbers here."]), "fr")
        assert not any(summary["fixes"].values())
    
    def test_invalid_target_language_raises(self, make_docx):
        with pytest.raises(ValueError):
            localize_document_bytes(make_docx(["text"]), "de")


class TestLocalizeDocument:
    def test_writes_a_localized_file_and_report(self, tmp_path, make_docx):
        path = tmp_path / "doc.docx"
        path.write_bytes(make_docx(["Value: 12.5%"]))
        output_path = localize_document(path, "fr")
        assert output_path == tmp_path / "doc_localized.docx"
        assert output_path.exists()
        assert _paragraph(output_path.read_bytes()).text == f"Value{NBSP}: 12,5{NBSP}%"
        report_path = tmp_path / "doc_localized_report.txt"
        assert report_path.exists()
        assert "Numbers reformatted" in report_path.read_text(encoding="utf-8")


# Cells that are purely numeric/punctuation — the only ones the localizer is
# responsible for. Acronym cells (ZINB, n.s., NF, ΔBIC) need glossary
# translation and are intentionally outside this test's scope.
_NUMERIC_CELL = re.compile(r"^[\d.,%/  -]+$")


class TestTableReplacementFixtures:
    @pytest.mark.parametrize("source_lang, target_lang", [("en", "fr"), ("fr", "en")])
    def test_numeric_cells_match_the_goal_fixture(self, source_lang, target_lang):
        source = Document(FIXTURE_DIR / f"test_table_replacements_{source_lang}.docx")
        goal = Document(FIXTURE_DIR / f"test_table_replacements_{target_lang}.docx")
        
        checked = 0
        for row, goal_row in zip(source.tables[0].rows, goal.tables[0].rows):
            for cell, goal_cell in zip(row.cells, goal_row.cells):
                if not _NUMERIC_CELL.match(cell.text):
                    continue
                localized = localize_text(cell.text, target_lang)[0]
                assert localized == goal_cell.text, (
                    f"{source_lang}->{target_lang}: {cell.text!r} localized to "
                    f"{localized!r}, expected {goal_cell.text!r}"
                )
                checked += 1
        assert checked > 0, "No numeric cells were found in the fixture"
    
    @pytest.mark.parametrize("acronym", ["ZINB", "n.s.", "NF", "ΔBIC"])
    def test_acronym_cells_are_left_untouched(self, acronym):
        # The localizer does not translate — acronym conversion is a glossary
        # concern handled elsewhere.
        assert localize_text(acronym, "fr")[0] == acronym
        assert localize_text(acronym, "en")[0] == acronym
