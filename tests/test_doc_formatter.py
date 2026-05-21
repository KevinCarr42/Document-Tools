import io

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

from src.doc_formatter import (
    _is_text_run,
    _run_signature,
    _run_text,
    format_document,
    format_document_bytes,
)

SENTENCE = "The quick brown fox jumps over the lazy dog."

_PROP_TAGS = {
    "bold": "w:b",
    "italic": "w:i",
    "underline": "w:u",
    "superscript": "w:vertAlign",
    "subscript": "w:vertAlign",
}


def _docx(paragraphs):
    # Each paragraph is a list of run fragments. A fragment is a plain string,
    # or a (text, props) tuple where props can hold bold/italic/underline/
    # superscript/subscript booleans and a "color" RGBColor.
    doc = Document()
    for fragments in paragraphs:
        paragraph = doc.add_paragraph()
        for fragment in fragments:
            text, props = fragment if isinstance(fragment, tuple) else (fragment, {})
            run = paragraph.add_run(text)
            if props.get("bold"):
                run.bold = True
            if props.get("italic"):
                run.italic = True
            if props.get("underline"):
                run.underline = True
            if props.get("superscript"):
                run.font.superscript = True
            if props.get("subscript"):
                run.font.subscript = True
            if "color" in props:
                run.font.color.rgb = props["color"]
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _split_at(text, indices):
    bounds = [0, *indices, len(text)]
    return [text[bounds[i]:bounds[i + 1]] for i in range(len(bounds) - 1)]


def _format(docx_bytes):
    # The cleaned document bytes only — drops the change summary.
    return format_document_bytes(docx_bytes)[0]


def _paragraphs(docx_bytes):
    return Document(io.BytesIO(docx_bytes)).paragraphs


def _run_count(paragraph):
    return len(paragraph._p.findall(qn("w:r")))


def _any_run_has_colour(paragraph):
    for run in paragraph._p.findall(qn("w:r")):
        rPr = run.find(qn("w:rPr"))
        if rPr is not None and rPr.find(qn("w:color")) is not None:
            return True
    return False


class TestFunnySplits:
    # Word fragments a paragraph into many runs — mid-word, mid-sentence, around
    # the undo stack of an autocorrect. Every split below is a single logical
    # sentence that must come back out whole, in one run, with spacing intact.
    @pytest.mark.parametrize(
        "indices",
        [
            [2],
            [3],
            [4],
            [9, 10],
            [3, 4],
            [5, 11, 18, 27, 33],
            list(range(1, len(SENTENCE))),
        ],
        ids=[
            "inside-first-word",
            "after-a-word",
            "start-of-a-word",
            "around-a-space",
            "space-alone-in-its-run",
            "several-mixed-breaks",
            "every-character-its-own-run",
        ],
    )
    def test_splits_recombine_into_one_run(self, indices):
        data = _docx([_split_at(SENTENCE, indices)])
        paragraphs = _paragraphs(_format(data))
        assert len(paragraphs) == 1
        assert _run_count(paragraphs[0]) == 1
        assert paragraphs[0].text == SENTENCE


class TestDisjointedRuns:
    def test_word_broken_in_the_middle(self):
        data = _docx([["wo", "rd"]])
        paragraph = _paragraphs(_format(data))[0]
        assert _run_count(paragraph) == 1
        assert paragraph.text == "word"
    
    def test_split_in_the_middle_of_a_sentence(self):
        data = _docx([["This is ", "a single ", "sentence."]])
        paragraph = _paragraphs(_format(data))[0]
        assert _run_count(paragraph) == 1
        assert paragraph.text == "This is a single sentence."
    
    def test_split_at_the_end_of_a_sentence(self):
        data = _docx([["First sentence. ", "Second sentence. ", "Third."]])
        paragraph = _paragraphs(_format(data))[0]
        assert _run_count(paragraph) == 1
        assert paragraph.text == "First sentence. Second sentence. Third."
    
    def test_many_funny_splits_in_one_paragraph(self):
        data = _docx([["Th", "e qu", "ick br", "own ", "fox", " jum", "ps."]])
        paragraph = _paragraphs(_format(data))[0]
        assert _run_count(paragraph) == 1
        assert paragraph.text == "The quick brown fox jumps."
    
    def test_empty_runs_are_dropped(self):
        data = _docx([["before", "", "", "after"]])
        paragraph = _paragraphs(_format(data))[0]
        assert _run_count(paragraph) == 1
        assert paragraph.text == "beforeafter"
    
    def test_already_clean_paragraph_stays_one_run(self):
        data = _docx([["Just one run."]])
        paragraph = _paragraphs(_format(data))[0]
        assert _run_count(paragraph) == 1
        assert paragraph.text == "Just one run."
    
    def test_each_paragraph_is_merged_independently(self):
        data = _docx([["one ", "two"], ["three ", "four"]])
        paragraphs = _paragraphs(_format(data))
        assert [p.text for p in paragraphs] == ["one two", "three four"]
        assert all(_run_count(p) == 1 for p in paragraphs)


class TestSpacePreservation:
    def test_space_attached_to_left_fragment(self):
        data = _docx([["hello ", "world"]])
        assert _paragraphs(_format(data))[0].text == "hello world"
    
    def test_space_attached_to_right_fragment(self):
        data = _docx([["hello", " world"]])
        assert _paragraphs(_format(data))[0].text == "hello world"
    
    def test_space_alone_in_its_own_run(self):
        data = _docx([["hello", " ", "world"]])
        paragraph = _paragraphs(_format(data))[0]
        assert _run_count(paragraph) == 1
        assert paragraph.text == "hello world"
    
    def test_double_space_between_words_is_kept(self):
        data = _docx([["a", "  ", "b"]])
        assert _paragraphs(_format(data))[0].text == "a  b"
    
    def test_long_run_of_spaces_is_not_collapsed(self):
        data = _docx([["word", "     ", "word"]])
        paragraph = _paragraphs(_format(data))[0]
        assert _run_count(paragraph) == 1
        assert paragraph.text == "word     word"
    
    def test_leading_and_trailing_spaces_are_kept(self):
        data = _docx([["  ", "padded", "  "]])
        paragraph = _paragraphs(_format(data))[0]
        assert _run_count(paragraph) == 1
        assert paragraph.text == "  padded  "
    
    def test_no_space_is_invented_between_mid_word_fragments(self):
        data = _docx([["super", "cali", "fragilistic"]])
        assert _paragraphs(_format(data))[0].text == "supercalifragilistic"
    
    def test_tab_between_runs_is_preserved(self):
        # A tab makes its run non-mergeable, but the text must survive intact.
        data = _docx([["left", "\t", "right"]])
        assert _paragraphs(_format(data))[0].text == "left\tright"


class TestManualColours:
    def test_manual_colour_is_removed(self):
        data = _docx([[("coloured text", {"color": RGBColor(0xFF, 0x00, 0x00)})]])
        paragraph = _paragraphs(_format(data))[0]
        assert not _any_run_has_colour(paragraph)
        assert paragraph.text == "coloured text"
    
    def test_recoloured_runs_merge_once_colour_is_reset(self):
        red = RGBColor(0xFF, 0x00, 0x00)
        data = _docx([[("red ", {"color": red}), ("text", {"color": red})]])
        paragraph = _paragraphs(_format(data))[0]
        assert _run_count(paragraph) == 1
        assert paragraph.text == "red text"
        assert not _any_run_has_colour(paragraph)


class TestAllowedFormatting:
    def test_bold_run_is_not_merged_with_a_plain_run(self):
        data = _docx([[("Bold", {"bold": True}), " plain"]])
        paragraph = _paragraphs(_format(data))[0]
        assert _run_count(paragraph) == 2
        assert paragraph.text == "Bold plain"
        assert paragraph.runs[0].bold is True
        assert paragraph.runs[1].bold is not True
    
    def test_adjacent_bold_runs_merge(self):
        data = _docx([[("Bo", {"bold": True}), ("ld", {"bold": True})]])
        paragraph = _paragraphs(_format(data))[0]
        assert _run_count(paragraph) == 1
        assert paragraph.text == "Bold"
        assert paragraph.runs[0].bold is True
    
    @pytest.mark.parametrize("prop", ["bold", "italic", "underline", "superscript", "subscript"])
    def test_allowed_formatting_survives_on_visible_text(self, prop):
        data = _docx([["plain ", ("marked", {prop: True})]])
        paragraph = _paragraphs(_format(data))[0]
        assert paragraph.text == "plain marked"
        assert _run_count(paragraph) == 2
        marked = paragraph._p.findall(qn("w:r"))[1]
        assert marked.find(qn("w:rPr")).find(qn(_PROP_TAGS[prop])) is not None


class TestWhitespaceOnlyRuns:
    def test_bold_is_stripped_from_a_space_only_run(self):
        # The space carried bold, which is not allowed; once stripped the run
        # collapses into its plain neighbours.
        data = _docx([["start", ("   ", {"bold": True}), "end"]])
        paragraph = _paragraphs(_format(data))[0]
        assert _run_count(paragraph) == 1
        assert paragraph.text == "start   end"
    
    def test_formatting_is_stripped_from_an_empty_run(self):
        data = _docx([["text", ("", {"italic": True, "underline": True})]])
        paragraph = _paragraphs(_format(data))[0]
        assert _run_count(paragraph) == 1
        assert paragraph.text == "text"
    
    def test_bold_space_between_bold_words_stays_separate(self):
        # The space cannot itself be bold, so it cannot fold into the words.
        data = _docx([[
            ("hello", {"bold": True}),
            ("   ", {"bold": True}),
            ("world", {"bold": True}),
        ]])
        paragraph = _paragraphs(_format(data))[0]
        assert paragraph.text == "hello   world"
        assert _run_count(paragraph) == 3
        assert paragraph.runs[1].text == "   "
        assert paragraph.runs[1].bold is not True


class TestSmartTags:
    def _docx_with_smart_tag(self, fragments):
        doc = Document()
        paragraph = doc.add_paragraph()
        smart_tag = OxmlElement("w:smartTag")
        for fragment in fragments:
            run = OxmlElement("w:r")
            text = OxmlElement("w:t")
            text.text = fragment
            text.set(qn("xml:space"), "preserve")
            run.append(text)
            smart_tag.append(run)
        paragraph._p.append(smart_tag)
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
    
    def test_smart_tag_runs_are_unwrapped_and_merged(self):
        data = self._docx_with_smart_tag(["New ", "York"])
        paragraph = _paragraphs(_format(data))[0]
        assert paragraph._p.find(qn("w:smartTag")) is None
        assert _run_count(paragraph) == 1
        assert paragraph.text == "New York"
    
    def test_smart_tag_run_merges_with_surrounding_runs(self):
        doc = Document()
        paragraph = doc.add_paragraph()
        paragraph.add_run("visit ")
        smart_tag = OxmlElement("w:smartTag")
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = "Paris"
        run.append(text)
        smart_tag.append(run)
        paragraph._p.append(smart_tag)
        paragraph.add_run(" today")
        buffer = io.BytesIO()
        doc.save(buffer)
        result = _paragraphs(_format(buffer.getvalue()))[0]
        assert result._p.find(qn("w:smartTag")) is None
        assert _run_count(result) == 1
        assert result.text == "visit Paris today"


class TestProofErrors:
    def test_proof_error_markers_are_removed(self):
        doc = Document()
        paragraph = doc.add_paragraph()
        paragraph.add_run("some")
        proof_error = OxmlElement("w:proofErr")
        proof_error.set(qn("w:type"), "spellStart")
        paragraph._p.append(proof_error)
        paragraph.add_run(" text")
        buffer = io.BytesIO()
        doc.save(buffer)
        result = _paragraphs(_format(buffer.getvalue()))[0]
        assert result._p.findall(qn("w:proofErr")) == []
        assert _run_count(result) == 1
        assert result.text == "some text"


class TestOrphanFieldRuns:
    def _field_run(self, fld_type):
        run = OxmlElement("w:r")
        fld_char = OxmlElement("w:fldChar")
        fld_char.set(qn("w:fldCharType"), fld_type)
        run.append(fld_char)
        return run
    
    def _field_runs(self, paragraph):
        return [r for r in paragraph._p.findall(qn("w:r")) if r.find(qn("w:fldChar")) is not None]
    
    def test_orphan_field_end_run_is_removed(self):
        doc = Document()
        paragraph = doc.add_paragraph()
        paragraph.add_run("text")
        paragraph._p.append(self._field_run("end"))
        buffer = io.BytesIO()
        doc.save(buffer)
        result = _paragraphs(_format(buffer.getvalue()))[0]
        assert self._field_runs(result) == []
        assert result.text == "text"
    
    def test_balanced_field_runs_are_kept(self):
        doc = Document()
        paragraph = doc.add_paragraph()
        paragraph._p.append(self._field_run("begin"))
        paragraph._p.append(self._field_run("end"))
        buffer = io.BytesIO()
        doc.save(buffer)
        result = _paragraphs(_format(buffer.getvalue()))[0]
        assert len(self._field_runs(result)) == 2
    
    def test_field_spanning_paragraphs_is_kept(self):
        # A field whose begin and end sit in different paragraphs must survive:
        # begin/end are balanced across the whole story, not one paragraph.
        doc = Document()
        doc.add_paragraph()._p.append(self._field_run("begin"))
        doc.add_paragraph()._p.append(self._field_run("end"))
        buffer = io.BytesIO()
        doc.save(buffer)
        paragraphs = _paragraphs(_format(buffer.getvalue()))
        assert len(self._field_runs(paragraphs[0])) == 1
        assert len(self._field_runs(paragraphs[1])) == 1


class TestStoryCoverage:
    def test_runs_in_a_table_cell_are_merged(self):
        doc = Document()
        cell = doc.add_table(rows=1, cols=1).cell(0, 0)
        cell.paragraphs[0].add_run("cell ")
        cell.paragraphs[0].add_run("text")
        buffer = io.BytesIO()
        doc.save(buffer)
        out = _format(buffer.getvalue())
        cell_paragraph = Document(io.BytesIO(out)).tables[0].cell(0, 0).paragraphs[0]
        assert _run_count(cell_paragraph) == 1
        assert cell_paragraph.text == "cell text"
    
    def test_runs_in_a_header_are_merged(self):
        doc = Document()
        header_paragraph = doc.sections[0].header.paragraphs[0]
        header_paragraph.add_run("header ")
        header_paragraph.add_run("text")
        buffer = io.BytesIO()
        doc.save(buffer)
        out = _format(buffer.getvalue())
        result = Document(io.BytesIO(out)).sections[0].header.paragraphs[0]
        assert _run_count(result) == 1
        assert result.text == "header text"
    
    def test_runs_inside_a_hyperlink_are_merged(self):
        doc = Document()
        paragraph = doc.add_paragraph()
        hyperlink = OxmlElement("w:hyperlink")
        for fragment in ["click ", "here"]:
            run = OxmlElement("w:r")
            text = OxmlElement("w:t")
            text.text = fragment
            text.set(qn("xml:space"), "preserve")
            run.append(text)
            hyperlink.append(run)
        paragraph._p.append(hyperlink)
        buffer = io.BytesIO()
        doc.save(buffer)
        out = _format(buffer.getvalue())
        link = Document(io.BytesIO(out)).paragraphs[0]._p.find(qn("w:hyperlink"))
        link_runs = link.findall(qn("w:r"))
        assert len(link_runs) == 1
        assert "".join(t.text for t in link_runs[0].findall(qn("w:t"))) == "click here"


class TestChangeSummary:
    def test_summary_counts_merged_runs(self):
        _, summary = format_document_bytes(_docx([["a ", "b ", "c"]]))
        assert summary["fixes"]["merged_runs"] == 2
        assert summary["paragraphs_scanned"] == 1
        assert summary["locations"]["body"] == 1
    
    def test_summary_counts_manual_colours(self):
        data = _docx([[("red", {"color": RGBColor(0xFF, 0x00, 0x00)})]])
        _, summary = format_document_bytes(data)
        assert summary["fixes"]["manual_colours"] == 1
    
    def test_summary_classifies_table_and_body_paragraphs(self):
        doc = Document()
        doc.add_paragraph("body text")
        doc.add_table(rows=1, cols=1).cell(0, 0).paragraphs[0].add_run("cell text")
        buffer = io.BytesIO()
        doc.save(buffer)
        _, summary = format_document_bytes(buffer.getvalue())
        assert summary["locations"]["body"] >= 1
        assert summary["locations"]["tables"] >= 1
    
    def test_clean_document_reports_no_fixes(self):
        _, summary = format_document_bytes(_docx([["one clean run"]]))
        assert not any(summary["fixes"].values())


class TestFormatDocument:
    def test_writes_a_formatted_file_next_to_the_input(self, tmp_path):
        path = tmp_path / "messy.docx"
        path.write_bytes(_docx([["one ", "run ", "split"]]))
        output_path = format_document(path)
        assert output_path == tmp_path / "messy_formatted.docx"
        assert output_path.exists()
        paragraph = _paragraphs(output_path.read_bytes())[0]
        assert _run_count(paragraph) == 1
        assert paragraph.text == "one run split"
    
    def test_writes_a_change_report(self, tmp_path):
        path = tmp_path / "messy.docx"
        path.write_bytes(_docx([["one ", "run ", "split"]]))
        format_document(path)
        report_path = tmp_path / "messy_formatted_report.txt"
        assert report_path.exists()
        assert "Disjointed runs merged" in report_path.read_text(encoding="utf-8")


class TestRunHelpers:
    def test_run_text_concatenates_text(self):
        run = _paragraphs(_docx([["alpha"]]))[0]._p.find(qn("w:r"))
        assert _run_text(run) == "alpha"
    
    def test_text_only_run_is_a_text_run(self):
        run = _paragraphs(_docx([["plain"]]))[0]._p.find(qn("w:r"))
        assert _is_text_run(run) is True
    
    def test_run_with_a_tab_is_not_a_text_run(self):
        run = _paragraphs(_docx([["\t"]]))[0]._p.find(qn("w:r"))
        assert _is_text_run(run) is False
    
    def test_identically_formatted_runs_share_a_signature(self):
        runs = _paragraphs(_docx([[("a", {"bold": True}), ("b", {"bold": True})]]))[0]._p.findall(qn("w:r"))
        assert _run_signature(runs[0]) == _run_signature(runs[1])
    
    def test_differently_formatted_runs_have_different_signatures(self):
        runs = _paragraphs(_docx([[("a", {"bold": True}), "b"]]))[0]._p.findall(qn("w:r"))
        assert _run_signature(runs[0]) != _run_signature(runs[1])
