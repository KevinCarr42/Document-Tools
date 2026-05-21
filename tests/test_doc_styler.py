import io
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from src.doc_styler import apply_format_spec_bytes, available_templates

_FIXTURES = Path(__file__).parent / "fixtures"


def _styles_by_id(docx_bytes):
    return {s.style_id: s for s in Document(io.BytesIO(docx_bytes)).styles}


@pytest.fixture
def plain_docx():
    doc = Document()
    doc.add_heading("A heading", level=1)
    doc.add_paragraph("Some body text.", style="Body Text")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


@pytest.fixture
def csas_base():
    # A document built from Word's built-in themed styles: its headings carry
    # borders and shading, and it uses a non-CSAS theme — all of which the spec
    # must override.
    return (_FIXTURES / "csas_example.docx").read_bytes()


class TestTemplates:
    def test_available_templates(self):
        templates = available_templates()
        assert "advisory_response_report" in templates
        assert "research_document" in templates
    
    def test_unknown_template_raises(self, plain_docx):
        with pytest.raises(KeyError):
            apply_format_spec_bytes(plain_docx, "no_such_template")


class TestApplySpec:
    def test_returns_bytes_and_summary(self, plain_docx):
        result, summary = apply_format_spec_bytes(
            plain_docx, "advisory_response_report")
        assert isinstance(result, bytes)
        assert summary["template"] == "CSAS Advisory / Response Report"
        assert summary["styles_created"] > 0
        assert summary["styles_updated"] > 0
        assert summary["sections_adjusted"] == 1
    
    def test_output_reopens(self, plain_docx):
        result, _ = apply_format_spec_bytes(plain_docx, "research_document")
        Document(io.BytesIO(result))
    
    def test_body_text_is_arial_11(self, plain_docx):
        for template in available_templates():
            result, _ = apply_format_spec_bytes(plain_docx, template)
            body = _styles_by_id(result)["BodyText"]
            assert body.font.name == "Arial"
            assert body.font.size == Pt(11)
    
    def test_heading_scale_differs_between_templates(self, plain_docx):
        advisory, _ = apply_format_spec_bytes(
            plain_docx, "advisory_response_report")
        research, _ = apply_format_spec_bytes(plain_docx, "research_document")
        assert _styles_by_id(advisory)["Heading1"].font.size == Pt(16)
        assert _styles_by_id(research)["Heading1"].font.size == Pt(12)
    
    def test_custom_style_is_created(self, plain_docx):
        result, _ = apply_format_spec_bytes(
            plain_docx, "advisory_response_report")
        styles = _styles_by_id(result)
        assert "Caption-Figure" in styles
        assert styles["Caption-Figure"].font.italic is True
    
    def test_citation_has_hanging_indent(self, plain_docx):
        result, _ = apply_format_spec_bytes(
            plain_docx, "advisory_response_report")
        citation = _styles_by_id(result)["citation"]
        # A hanging indent is stored as a negative first-line indent.
        assert citation.paragraph_format.first_line_indent == Pt(-18)
    
    def test_hyperlink_keeps_spec_colour(self, plain_docx):
        result, _ = apply_format_spec_bytes(
            plain_docx, "advisory_response_report")
        hyperlink = _styles_by_id(result)["Hyperlink"]
        assert hyperlink.font.color.rgb == RGBColor(0x00, 0x00, 0xFF)
    
    def test_heading_colour_normalised_to_automatic(self, plain_docx):
        result, _ = apply_format_spec_bytes(plain_docx, "research_document")
        heading = _styles_by_id(result)["Heading1"]
        assert heading.font.color.rgb is None
    
    def test_page_setup_is_letter_one_inch_margins(self, plain_docx):
        result, _ = apply_format_spec_bytes(
            plain_docx, "advisory_response_report")
        section = Document(io.BytesIO(result)).sections[0]
        assert section.page_width == Pt(612)
        assert section.page_height == Pt(792)
        assert section.top_margin == Pt(72)
        assert section.left_margin == Pt(72)


class TestBordersShadingTheme:
    def test_heading_borders_and_shading_cleared(self, csas_base):
        result, _ = apply_format_spec_bytes(
            csas_base, "advisory_response_report")
        styles = _styles_by_id(result)
        for level in range(1, 6):
            pPr = styles[f"Heading{level}"].element.find(qn("w:pPr"))
            assert pPr.find(qn("w:pBdr")) is None
            assert pPr.find(qn("w:shd")) is None
    
    def test_cover_header_border_preserved(self, csas_base):
        result, _ = apply_format_spec_bytes(
            csas_base, "advisory_response_report")
        pPr = _styles_by_id(result)["CoverPageHeaderregions"].element.find(qn("w:pPr"))
        pBdr = pPr.find(qn("w:pBdr"))
        assert pBdr is not None
        assert pBdr.find(qn("w:bottom")) is not None
    
    def test_linked_char_style_stripped_of_built_in_decoration(self, csas_base):
        result, _ = apply_format_spec_bytes(
            csas_base, "advisory_response_report")
        rPr = _styles_by_id(result)["Heading1Char"].element.find(qn("w:rPr"))
        assert rPr.find(qn("w:shd")) is None
        assert rPr.find(qn("w:color")) is None
    
    def test_theme_is_replaced_with_csas_theme(self, csas_base):
        result, _ = apply_format_spec_bytes(
            csas_base, "advisory_response_report")
        theme = zipfile.ZipFile(io.BytesIO(result)).read("word/theme/theme1.xml")
        assert b"4F81BD" in theme
        assert b"Wood Type" not in theme
